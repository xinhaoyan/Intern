import os
import openai
import json
import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.enum.text import WD_COLOR_INDEX



openai.api_key = "EMPTY"
openai.api_base = "http://172.16.16.19X:XXXX/v1"

def get_completion(message, model="/data/models/01-Yi/Yi-34B-Chat"):
    response = openai.ChatCompletion.create(
        model=model,
        messages=message,
        max_tokens=1000,
        top_p=0.8,
        temperature=0.1,
        stop_token_ids=[7],
        stream=True,
    )
    full_message = ""
    for chunk in response:
        if "content" in chunk.choices[0].delta.keys() and chunk.choices[0].delta.content is not None:
            full_message += chunk.choices[0].delta.content
            
    
    return full_message

# def process_file_content(file_content):
#           return process_chunk(file_content)


def process_chunk(chunk):
    #创建包含系统和用户信息的message列表
    message = [
    {
      "role": "system",
      "content": "任务说明：\n作为小说网站的管理员，你的任务是识别并提取网络流行小说中的广告内容。请根据以下步骤和标准来分析文本，并确定是否存在广告：\n\n- 广告只会位于小说的开头或结尾，并可能以特定标志如'PS:'，'【'，'（'，'——'等开头。\n\n- 广告内容主要是小说作者为了推广而与读者互动的部分，与小说情节发展无关。互动语气包括但不限于“祈求”，“评论”，“批评”，“祝福”。“问候”。\n\n- 广告一般只有一句话，简洁明了。但是偶尔小说中广告会出现多句话，如果你判断小说中的广告为多句话为广告，请你谨慎判断，确保不会把小说内容识别为广告词。\n\n- 广告可能包括作者推荐小说的句子、打赏要求，或对小说内容的说明、解释和评价。\n\n- 注意，大部分小说内是没有广告的，请不要过度解读，从而把小说内容误解成广告。\n\n请在分析小说内容时，特别注意区分广告内容和正文，重点关注以特定标志开始的段落和与小说情节无关的简短互动性语句。\n\n---\n\n请你按照以下的步骤来找出广告：\n\n1.阅读一段连载小说的文本，注意其叙述风格和结构。\n\n2.识别文本中可能的风格变化和与主故事线不一致的元素。\n\n3.特别关注文末内容，判断是否有作者直接与读者互动的部分，或者是否包含广告语言和非故事元素。\n\n4. 分析这些元素是否构成广告或仅仅是作者试图增加读者互动的方式。\n\n【以下为小说内容：】\n### 您的小说文本 ###\n---\n在分析后，模型应按照JSON的格式输出：\nKey  为\\\\\\\"广告内容\\\\\\\"， Value 为识别到的广告内容或者\\\\\\\"无广告\\\""
    },
    {
      "role": "user",
      "content": """ <小说文本示例1>:'''过了一会男人又发了信息过来，问顾念下榻的酒店叫什么。\n    顾念还真的没注意这个，她翻了翻子豪发给她的信息。\n    子豪给她订的酒店，据说正好沿海，开窗就能看见不远处的海边。\n    关于下榻酒店的事情，她就不太想和这男人说了。\n    为了防止男人再纠缠，她说自己有朋友在三亚，说朋友会过来接她。\n    回了信息之后，顾念就把手机放下了。\n    眼角有些不自觉的又瞄了一下池遇。\n    池遇一直没反应，好像是对她这边的动静一点也不关心。\n    顾念突然就觉得有些无趣。\n    想起之前两个人在婚姻存续期间内，自己出门被人搭讪，池遇知道后都没什么反应。\n    如今这样，也算是正常了。\n    因为不爱，所以能不被影响。\n    这个认知，真的是让顾念心里不舒服的很。\n    在之前将近一年的时间内，她不是没努力过。\n    只是池遇真的是一块难啃的骨头。\n PS：请大家看清楚了，顾念可不是好人啊'''"""
    },
    {
      "role": "assistant",
      "content":  """{"广告内容":"PS：请大家看清楚了，顾念可不是好人啊"}"""
    }, 
    {
      "role": "user",
      "content": """<小说文本示例2>:'''PS: 最近太累啦，随便写写\n————\n先行是近几年在超一线城市迅速崛起并被富人极力追捧的教育机构，业内称其为天才集中营，外界则戏称烧钱大本营，诉其集嫌贫爱富和攀高结贵于一体。\n    虽风评天差地别，但没人可以否认先行在教育行业的突出表现。先行主要服务面临中高考的学生，承诺无论基础如何，保证最迟一年时间考上满意院校。\n    创办整三年，承诺百分百兑现，引来大批望子成龙的富人拿着钱排队预约，当普通人还在花着大把时间拼着不确定的未来时，有些人早已拿钱买注定灿烂的未来了。\n    早上八点多，闵姜西出现在cbd最豪华地段，买了早餐和牛奶，预留出跟几十上百号人争抢电梯的时间，来到公司的时候，距离正式上班还有十几分钟。\n    往常的清晨总是最百无聊赖的时刻，即便早到的人也都坐在各自的座位上，或对着镜子补妆，或对着电脑补课，静得像是临近高考的实验班，然而今天情况很是特殊，闵姜西一推门便看到一帮人聚在一起，似乎发生了什么大事儿。'''"""
    },
    {
      "role": "assistant",
      "content": """{"广告内容":"PS: 最近太累啦，随便写写"}"""
    },
    {
        "role": "user",
        "content": """ <小说文本示例3>:'''闵姜西看着不动声色，实则心底警铃大作，她突然想到昨晚她在车上，秦佔发给她的那条短信，陆遇迟装警察的事情，原本只有他们两个知道，可秦佔却神不知鬼不觉的发现了。\n    她不确定秦佔是什么时候知晓的，是只查了陆遇迟一个人，还是像外界传言的那般，秦佔所在的方圆千米内，不可能有身份可疑的人，就怕对他图谋不轨。\n    不管是碰巧还是意料之中，越是跟秦佔接触，闵姜西就越觉着传言非虚，怪不得程双光是听到他的名字就如临大敌。\n    短暂的如鲠在喉，闵姜西很快便强迫自己镇定对应:“谢谢秦先生给我机会，我会努力成为自己人。”\n    秦佔抽了最后一口烟，将烟头按灭在一旁的灭烟器中，淡淡道:“明天上午十点。”\n    闵姜西点头应声，“好。”'''"""
    },
    {
        "role": "assistant",
        "content": """{"广告内容":"无广告"}"""
    },
    {
        "role": "user",
        "content": f"""<小说文本示例4>: '''{chunk}'''"""
    }
        ]
    return {"result": get_completion(message, model="/data/models/01-Yi/Yi-34B-Chat"), "chunk": chunk}
    # return {"chunk": chunk}
    

# 函数：按句子分割文本
def split_into_chunks_by_sentences(text, max_chars=1000):
    sentences = text.split('\n')
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) + 1 > max_chars:
            chunks.append((current_chunk, len(current_chunk)))
            current_chunk = sentence + "\n"
        else:
            current_chunk += sentence + "\n"

    if current_chunk:
        chunks.append((current_chunk, len(current_chunk)))

    return chunks

# 修改后的 handle_large_file 函数
def handle_large_file(file_content):
    chunks = split_into_chunks_by_sentences(file_content, max_chars=1000)
    results = []
    original_lengths = []  # 存储原始chunk的长度
    previous_chunk_content = ""

    for chunk, chunk_length in chunks:
        # 存储原始chunk的长度
        original_lengths.append(chunk_length)

        # 检查chunk的长度
        if chunk_length < 500:
            # 从上一个chunk中抽取内容
            additional_content = previous_chunk_content[-500:]
            chunk = additional_content + chunk
        elif chunk.count('\n') == 1:  # 只有一句话
            additional_content = previous_chunk_content[-700:]
            chunk = additional_content + chunk

        try:
            chunk_data = process_chunk(chunk)
            results.append(chunk_data)
        except Exception as e:
            print(f"处理块时发生错误: {e}")
            continue

        # 更新previous_chunk_content
        previous_chunk_content = chunk

    return results, original_lengths


def sorted_filenames(folder_path):
    def extract_number(filename):
        match = re.search(r'(\d+)', filename)
        return int(match.group(1)) if match else float('inf')

    return sorted(os.listdir(folder_path), key=extract_number)
def process_file_content(file_content, previous_sentences=None):
    if previous_sentences:
        file_content = previous_sentences + '\n' + file_content
    return handle_large_file(file_content)

def clean_json_string(s):
    """
    清理JSON字符串，移除或转义可能导致解析错误的字符。
    """
    s = s.replace('\n', '\\n').replace('\r', '\\r')
    return s

def process_folder(folder_path, output_file):
    results = []
    previous_sentences = None

    for filename in sorted_filenames(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path) and file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                file_content = file.read()
                sentences = file_content.strip().split('\n')

                if len(sentences) <= 5:
                    chunk_data, lengths = process_file_content(file_content, previous_sentences)
                else:
                    chunk_data, lengths = process_file_content(file_content)

                    # 更新previous_sentences为当前文件的前五句话
                    previous_sentences = '\n'.join(sentences[-5:])
                for i,chunk in enumerate(chunk_data):
                     if isinstance(chunk, dict) and "result" in chunk and "chunk" in chunk:
                        result_str  = chunk["result"]
                        chunk_content = chunk["chunk"] # 使用文件内容作为 chunk
                        cleaned_result  = clean_json_string(result_str)
                        try :
                            
                            first_brace = cleaned_result.find("{")
                            last_brace = cleaned_result.rfind("}")
                            if first_brace != -1 and last_brace != -1 and first_brace < last_brace:
                                json_string = cleaned_result[first_brace:last_brace + 1]
                                json_obj = json.loads(json_string)
                                result_data = {"filename": filename, "result": json_obj, "chunk":  chunk_content,"original_lengths": lengths[i]}
                                results.append(result_data)
                                print(f"Result for {filename}:\n{chunk_data}\n {lengths}\n")  # 打印正常存入JSON的内容
                        except json.JSONDecodeError as e:
                                error_message = f"JSON解析错误: {e}"
                                print(f"{error_message}, 文件：{filename}")
                                new_result = "{\"广告内容\":\"" + json_string.replace("\"", "\\\"") + "\"}"
                                result_data = {"filename": filename, "result": new_result, "chunk":  chunk_content,"original_lengths": lengths[i]}
                                results.append(result_data)
                                print(f"Changed Result for {filename}:\n{chunk_data}\n {lengths}\n")

    # 将结果保存到JSON文件
    if results:
        json_output = json.dumps(results, indent=4, ensure_ascii=False)
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(json_output)
        print(f"Results saved to {output_file}")
    else:
        print("No advertisements found in any files.")



# 文件夹列表
# folder_list = [ "Novel_1", "Novel_2", "Novel_3","Novel_4", "Novel_5", "Novel_6","Novel_7", "Novel_8", "Novel_9", "Novel_10","Novel_11", "Novel_12", "Novel_13","Novel_14", "Nove      l_15", "Novel_16","Novel_17", "Novel_18", "Novel_19", "Novel_20","Novel_21"]

# process_multiple_folders(folder_list)

# folder_path = "Novel_22"
# output_file = "Novel_22.json"
# process_folder(folder_path, output_file)    
# folder_path = "Novel_23"
# output_file = "Novel_23.json"
# process_folder(folder_path, output_file)  

def load_json_data(json_file):
    with open(json_file, 'r', encoding='utf-8') as file:
        data = json.load(file)
    return data

def longest_common_substring_length(str1, str2):
    """ 计算两个字符串的最长公共子串长度，并返回其在str1中的起始位置 """
    m, n = len(str1), len(str2)
    dp = [[0] * (n + 1) for _ in range(m + 1)]
    max_length, start_pos = 0, 0

    for i in range(1, m + 1):
        for j in range(1, n + 1):
            if str1[i - 1] == str2[j - 1]:
                dp[i][j] = dp[i - 1][j - 1] + 1
                if dp[i][j] > max_length:
                    max_length = dp[i][j]
                    start_pos = i - max_length

    return max_length, start_pos



def filter_ad_content_based_on_similarity(ad_content, chunk, original_length, threshold=0.5):
    """ 根据与chunk的非原始长度部分的相似度过滤广告内容 """
    non_original_part = chunk[:-original_length]
    lcs_length, start_idx = longest_common_substring_length(ad_content, non_original_part)
    similarity = lcs_length / len(ad_content) if len(ad_content) > 0 else 0 
    
    if similarity >= threshold:
        end_idx = start_idx + lcs_length
        return ad_content[:start_idx] + ad_content[end_idx:]
    return ad_content


def add_highlighted_text(paragraph, text, highlight_text, similarity_threshold=0.3):
    """ 根据最大公共子串的相似度决定是否高亮文本 """
    lcs_length, start_idx = longest_common_substring_length(text, highlight_text)
    similarity = lcs_length / len(highlight_text) if len(highlight_text) > 0 else 0 
    print(f"similarity:{similarity}")
    if similarity >= similarity_threshold:
        # 查找并高亮显示最长公共子串
        print(f"text:{text}")
        # start_idx_in_text = text.find(highlight_text[start_idx:start_idx + lcs_length])
        print(f"Start Index: {start_idx}, Highlight Text: {highlight_text[:lcs_length]},text :{text}")
        if start_idx != -1:
            end_idx = start_idx + lcs_length
            paragraph.add_run(text[:start_idx])
            highlighted_run = paragraph.add_run(text[start_idx:end_idx])
            highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            paragraph.add_run(text[end_idx:])
        else:
            paragraph.add_run(text)
    else:
        paragraph.add_run(text)
        
def create_single_docx_from_chunks_with_highlights(data, output_file, include_all_chapters=False):
    doc = Document()

    # 按filename分组数据
    chapters = {}
    for item in data:
        filename = item['filename']
        if filename not in chapters:
            chapters[filename] = []
        chapters[filename].append(item)

    # 存储满足条件的filename
    valid_filenames = set()

    # 预处理，确定哪些filename满足条件
    for filename, items in chapters.items():
        for item in items:
            result = item['result']
            if isinstance(result, str):
                result = json.loads(result)
            ad_content = result['广告内容']
            chunk = item['chunk']
            original_length = item['original_lengths']

            if ad_content != "无广告":
                filtered_ad_content = filter_ad_content_based_on_similarity(ad_content, chunk, original_length)
                if ad_content in chunk and len(filtered_ad_content) > 0:
                    valid_filenames.add(filename)
                    break

    last_filename = None

    # 遍历每个章节
    for filename, items in chapters.items():
        # 只处理满足条件的filename
        if filename not in valid_filenames and not include_all_chapters:
            continue

        if last_filename is not None and last_filename != filename:
            doc.add_paragraph("\n")

        for item in items:
            result = item['result']
            if isinstance(result, str):
                result = json.loads(result)

            chunk = item['chunk']
            original_length = item['original_lengths']
            content_to_add = chunk[-original_length:]
            ad_content = result['广告内容']

            if ad_content != "无广告":
                ad_content = filter_ad_content_based_on_similarity(ad_content, chunk, original_length)
                if len(ad_content) > 0:
                    add_highlighted_text(doc.add_paragraph(), content_to_add, ad_content)
                else:
                    doc.add_paragraph(content_to_add)
            else:
                doc.add_paragraph(content_to_add)

        last_filename = filename

    doc.save(output_file)



# folder_path = "Test2"
# output_file = "Test2.json"
# process_folder(folder_path, output_file)     
# json_file = 'Test2.json'  # 替换为您的JSON文件路径
# output_folder = 'Test2_Docx.docx'  # 替换为您希望保存docx文件的文件夹路径
# # 加载JSON数据并创建docx文档
# data = load_json_data(json_file)
# create_single_docx_from_chunks_with_highlights(data, output_folder)


# folder_path = "Novel_22"
# output_file = "Novel_22.json"
# process_folder(folder_path, output_file)     
# json_file = 'Novel_22.json'  # 替换为您的JSON文件路径
# output_folder = 'Novel_22_Docx.docx'  # 替换为您希望保存docx文件的文件夹路径
# # 加载JSON数据并创建docx文档
# data = load_json_data(json_file)
# create_single_docx_from_chunks_with_highlights(data, output_folder)




# folder_path = "Novel_23"
# output_file = "Novel_23.json"
# process_folder(folder_path, output_file) 
# # JSON文件路径
# json_file = 'Novel_23.json'  # 替换为您的JSON文件路径
# output_folder = 'Novel_23_Docx.docx'  # 替换为您希望保存docx文件的文件夹路径
# # 加载JSON数据并创建docx文档
# data = load_json_data(json_file)
# create_single_docx_from_chunks_with_highlights(data, output_folder)


# folder_path = "Novel_24"
# output_file = "Novel_24.json"
# process_folder(folder_path, output_file) 
# # JSON文件路径
# json_file = 'Novel_24.json'  # 替换为您的JSON文件路径
# output_folder = 'Novel_24_Docx.docx'  # 替换为您希望保存docx文件的文件夹路径
# # 加载JSON数据并创建docx文档
# data = load_json_data(json_file)
# create_single_docx_from_chunks_with_highlights(data, output_folder)

folder_path = "Novel_Dataset/Novel_4"
output_file = "Novel_4.json"
process_folder(folder_path, output_file) 
# JSON文件路径
json_file = 'Novel_4.json'  # 替换为您的JSON文件路径
output_folder = 'Novel_4_Docx.docx'  # 替换为您希望保存docx文件的文件夹路径
# 加载JSON数据并创建docx文档
data = load_json_data(json_file)
create_single_docx_from_chunks_with_highlights(data, output_folder)
